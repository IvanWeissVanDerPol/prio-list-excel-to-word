import os
from numpy import nan
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from math import floor

from sqlalchemy import table

#!####################################
#!####################################
# JIRA column order:
#! issue_ID Customer Description Summary FixBuild Labels Status 
#!####################################
#!####################################


driver_path = "driver\chromedriver.exe"
excel_path = "\\\\192.168.2.3\\Share\\PrioList\\Prio List.xlsx"
word_path = "\\\\192.168.2.3\\Share\\PrioList\\Build_Planning_Report.docx"
log_in_url  =  'http://192.168.1.3:8080/login.jsp?os_destination=%2Fsecure%2FMyJiraHome.jspa'
search_page_url_template = "http://192.168.1.3:8080/issues/?jql=%22Issue%20ID%22%20%3D%20"
ticket_page_url_template = 'http://192.168.1.3:8080/browse/'
to_buil_durl = "http://192.168.1.3:8080/issues/?jql=labels%20%3D%20TO_BUILD%20ORDER%20BY%20status%20DESC"

#Credentials for jira
user='label2build'
password='label2build'

chrome_options = Options()
chrome_options.add_experimental_option("detach", True)

try:
    os.remove(word_path)
except:
    pass

#make the document
my_doc = Document()
sections = my_doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
my_doc.add_heading('Build planning report', 0)


pd.set_option('display.max_columns', None)
#read the prio list excel and make an auxiliary dataframe
df_excel = pd.read_excel(excel_path,skiprows=1)
df_excel.fillna('', inplace=True)

#clean the initial dataframe and order the columns
df_excel = df_excel.drop(columns=["Tester", "Developer", "Action holder WPG", "Due Date", "Days to due date"])
df_excel = df_excel[["Branch","Prio Nr.","Customer", "Ticket Nr.", "Description", "Current status"]]
#add a column for the labels
df_excel['labels'] = ""
df_excel['Branch'] = df_excel['Branch'].str.rsplit('/', n=1).str.get(0)

#remove unwanted rows
list_of_strings_that_mean_ticket_should_be_removed = ["close"] # closed tickets
for string_to_check in list_of_strings_that_mean_ticket_should_be_removed:
    df_excel[~df_excel["Current status"].str.contains(string_to_check)==False]



#  the support dataframe
df_Support = df_excel[df_excel["Prio Nr."] == 0]
df_Support = df_Support.drop(columns=["Prio Nr."])
df_Support = df_Support[df_Support['Current status'].str.contains("upport")]

browser = webdriver.Chrome(executable_path=driver_path)
#log in to jira
browser.get(log_in_url)
browser.find_element_by_id('login-form-username').send_keys(user)
browser.find_element_by_id('login-form-password').send_keys(password)
browser.find_element_by_id('login-form-submit').click()

df_Support = df_Support.sort_values(["Customer", "Ticket Nr." ], ascending = (True, True))

my_doc.add_heading('Support Tickets:', 1)
paragraph = my_doc.add_paragraph()
for row in df_Support.itertuples():
    issue_ID = df_Support.loc[row.Index, "Ticket Nr."]
    customer = df_Support.loc[row.Index, "Customer"]
    description = df_Support.loc[row.Index, "Description"]
    paragraph.add_run("* " + issue_ID  + "\t-\t" + customer + "\t-\t" + description + "\n")


######  fixed  dataframe  ######
#open browser to the to build url and exctract data to a dictionary that i append to the dataframe as a row 
browser.get(to_buil_durl)
issuetable = browser.find_element_by_xpath('//*[@id="issuetable"]/tbody')
table_rows = issuetable.find_elements_by_xpath("./tr")
df_Fixed = pd.DataFrame()
df_Fixed["Ticket Nr."] = ""
df_Fixed["Customer"] = ""
df_Fixed["Current Status"] = ""
df_Fixed["Branch"] = ""
for row in table_rows:
    elements = row.find_elements_by_xpath("./td")
    matrix = {}
    matrix["Ticket Nr."] = elements[0].text
    matrix["Customer"] = elements[1].text
    matrix["Current Status"] = elements[3].text
    builds = elements[4].text
    branch_list = builds.replace(",", ";")
    for branch in branch_list.split(" "):   # repete each row for each branch it has
        matrix["Branch"] = branch
        df_Fixed = df_Fixed.append(matrix, ignore_index = True)

df_Fixed = df_Fixed.sort_values(["Branch", "Customer", "Ticket Nr." ], ascending = (True, True,True))

my_doc.add_heading('The next tickets are ready to deliver:', 1)
new_branch_found = True
new_customer_found = True
current_branch = ""
current_customer = ""
for row in df_Fixed.itertuples():
    prev_branch = current_branch
    current_branch =df_Fixed.loc[row.Index, "Branch"]
    if prev_branch != current_branch:
        new_branch_found = True
        paragraph = my_doc.add_heading(current_branch, level=2)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(0)
        current_customer = ""
        new_branch_found = False
    prev_customer = current_customer
    current_customer = df_Fixed.at[row.Index,'Customer']
    if str(current_customer) == "nan":
        current_customer = ""
    issue_ID = df_Fixed.loc[row.Index, "Ticket Nr."]
    Current_Status = df_Fixed.loc[row.Index, "Current Status"]
    paragraph = my_doc.add_paragraph("  *  " + issue_ID + "\t-\t" + current_customer + "\t-\t")
    paragraph.add_run(Current_Status)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)



######  pending dataframe  ######
#make the pending dataframe
df_pending = pd.DataFrame()
df_pending["Branch"] = ""
df_pending["Ticket Nr."] = ""
df_pending["Customer"] = ""
df_pending["Current Status"] = ""
df_pending["Description"] = ""

#extract the important info from the non prio 0 dataframe
df_prioNot0 = df_excel[df_excel["Prio Nr."] != 0]
for row in df_prioNot0.itertuples():
    prio = df_prioNot0.at[row.Index,'Prio Nr.']
    branches = df_prioNot0.at[row.Index,'Branch']
    customer = df_prioNot0.at[row.Index,'Customer']
    ticket_id = df_prioNot0.at[row.Index,'Ticket Nr.']
    description = df_prioNot0.at[row.Index,'Description']
    Current_Status = df_prioNot0.at[row.Index,'Current status']
    branch_list = branches.split("/")
    for banch in branch_list:
        aux_dic = {}
        aux_dic["Branch"] = banch
        aux_dic["Ticket Nr."] = ticket_id
        aux_dic["Customer"] = customer
        aux_dic["Description"] = description
        aux_dic["Current Status"] = Current_Status
        aux_dic["Prio Nr."] = prio
        df_pending =df_pending.append(aux_dic, ignore_index = True)
df_pending = df_pending.sort_values(["Branch", "Customer", "Ticket Nr." ], ascending = (True, True,True))

my_doc.add_heading('Other tickets status (Order by Customer / Ticket Nr.):', 1)
new_branch_found = True
new_customer_found = True
current_branch = ""
current_customer = ""
for row in df_pending.itertuples():
    prev_branch = current_branch
    current_branch =df_pending.loc[row.Index, "Branch"]
    if prev_branch != current_branch:
        new_branch_found = True
        paragraph = my_doc.add_heading(current_branch, level=2)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(0)
        current_customer = ""
        new_branch_found = False
    prev_customer = current_customer
    current_customer = df_pending.at[row.Index,'Customer']
    if str(current_customer) == "nan":
        current_customer = ""
    issue_ID = df_pending.loc[row.Index, "Ticket Nr."]
    Current_Status = df_pending.loc[row.Index, "Current Status"]
    prio = str(round(df_pending.loc[row.Index, "Prio Nr."]))
    description = df_pending.loc[row.Index, "Description"]
    Current_Status = df_pending.loc[row.Index, "Current Status"]
    paragraph = my_doc.add_paragraph("  *  " + issue_ID + "\t-\t" + current_customer + "\t-\t"+ prio + "\t-\t"+ description + "\t-\t")
    paragraph.add_run(Current_Status).bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

my_doc.save(word_path)